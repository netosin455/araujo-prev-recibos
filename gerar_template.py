#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_paragraph(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Helper para adicionar parágrafos formatados"""
    p = doc.add_paragraph(text)
    p.alignment = align
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    return p

def create_receipt_template():
    """Cria um template de recibo com campos dinâmicos"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ── CABEÇALHO ──
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    table.allow_autofit = False
    
    # Logo (coluna 1)
    cell_logo = table.rows[0].cells[0]
    cell_logo.width = Inches(1.2)
    p_logo = cell_logo.paragraphs[0]
    p_logo.text = "{{ empresa_logo }}"
    
    # Dados da empresa (coluna 2-3)
    cell_empresa = table.rows[0].cells[1]
    cell_empresa.merge(table.rows[0].cells[2])
    
    p_empresa = cell_empresa.paragraphs[0]
    p_empresa.text = "{{ empresa_nome }}"
    p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_empresa.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    
    p_cnpj = cell_empresa.add_paragraph("CNPJ: {{ empresa_cnpj }} | IE: {{ empresa_ie }}")
    p_cnpj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_cnpj.runs:
        run.font.size = Pt(9)
    
    p_endereco = cell_empresa.add_paragraph("{{ empresa_endereco }} | {{ empresa_cidade }}")
    p_endereco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_endereco.runs:
        run.font.size = Pt(9)
    
    p_contato = cell_empresa.add_paragraph("Tel: {{ empresa_telefone }} | Email: {{ empresa_email }}")
    p_contato.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_contato.runs:
        run.font.size = Pt(9)
    
    # Remover bordas da tabela
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # ── LINHA SEPARADORA ──
    doc.add_paragraph("_" * 80)
    
    # ── TÍTULO ──
    p_titulo = add_paragraph(doc, "RECIBO DE HONORÁRIOS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── NÚMERO DO RECIBO ──
    p_num = add_paragraph(doc, f"Nº {{ num_recibo }}", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── CORPO DO RECIBO ──
    doc.add_paragraph()  # Espaço
    
    p_recebemos = add_paragraph(doc, "Recebemos de:", size=11, bold=True)
    
    p_nome = add_paragraph(doc, "{{ nome }}", size=11)
    
    p_cpf = add_paragraph(doc, f"CPF/CNPJ: {{ cpf }}", size=11)
    
    p_cidade = add_paragraph(doc, f"Município/UF: {{ municipio_uf }}", size=11)
    
    doc.add_paragraph()  # Espaço
    
    p_valor_label = add_paragraph(doc, "A importância de:", size=11, bold=True)
    
    p_valor = add_paragraph(doc, f"R$ {{ valor }}", size=12, bold=True)
    
    doc.add_paragraph()  # Espaço
    
    p_descricao_label = add_paragraph(doc, "Referente a:", size=11, bold=True)
    
    p_descricao = add_paragraph(doc, "Ação Previdenciária{{ complemento_fmt }}", size=11)
    
    doc.add_paragraph()  # Espaço
    
    p_data_label = add_paragraph(doc, "Data:", size=11, bold=True)
    
    p_data = add_paragraph(doc, "{{ data_extenso }}", size=11)
    
    # ── ASSINATURA ──
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_assinatura = add_paragraph(doc, "_" * 40, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_assinatura_nome = add_paragraph(doc, "{{ empresa_nome }}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_assinatura_cpf = add_paragraph(doc, "CPF/CNPJ: {{ empresa_cnpj }}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Salvar template
    template_path = "template/recibo_template.docx"
    doc.save(template_path)
    print(f"✅ Template criado com sucesso em: {template_path}")

if __name__ == "__main__":
    create_receipt_template()
