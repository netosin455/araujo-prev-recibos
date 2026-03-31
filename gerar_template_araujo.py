#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_paragraph(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Helper para adicionar parágrafos formatados"""
    p = doc.add_paragraph(text)
    p.alignment = align
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    return p

def create_araujo_receipt_template():
    """Cria um template de recibo oficial da Araujo Prev"""
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # ── LOGO ──
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture("C:\\Users\\carlo\\Downloads\\Logo par forms.png", width=Cm(2))

    # ── CABEÇALHO OFICIAL ──
    p_titulo = add_paragraph(doc, "A ARAUJO SERVIÇOS LTDA ME", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_titulo.runs[0].font.color.rgb = RGBColor(30, 64, 175)
    
    p_subtitulo = add_paragraph(doc, "A ARAUJO PREV", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_cnpj = add_paragraph(doc, "CNPJ: {{ empresa_cnpj }} | IE: {{ empresa_ie }}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_endereco = add_paragraph(doc, "{{ empresa_endereco }}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_cidade = add_paragraph(doc, "{{ empresa_cidade }}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_contato = add_paragraph(doc, "Tel: {{ empresa_telefone }} | Email: {{ empresa_email }}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── LINHA SEPARADORA ──
    doc.add_paragraph("_" * 80)
    
    # ── TÍTULO DO DOCUMENTO ──
    p_titulo_doc = add_paragraph(doc, "RECIBO DE HONORÁRIOS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── NÚMERO DO RECIBO ──
    p_num = add_paragraph(doc, f"Nº {{ num_recibo }}", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── CORPO DO RECIBO ──
    doc.add_paragraph()  # Espaço
    
    p_recebemos = add_paragraph(doc, "Recebemos de:", size=11, bold=True)
    
    p_nome = add_paragraph(doc, "{{ nome }}", size=11)
    
    p_cpf = add_paragraph(doc, f"CPF/CNPJ: {{ cpf }}", size=11)
    
    p_cidade_cliente = add_paragraph(doc, f"Município/UF: {{ municipio_uf }}", size=11)
    
    doc.add_paragraph()  # Espaço
    
    p_valor_label = add_paragraph(doc, "A importância de:", size=11, bold=True)
    
    p_valor = add_paragraph(doc, f"R$ {{ valor }}", size=12, bold=True)
    
    doc.add_paragraph()  # Espaço
    
    p_descricao_label = add_paragraph(doc, "Referente a:", size=11, bold=True)
    
    p_descricao = add_paragraph(doc, "Ação Previdenciária{{ complemento_fmt }}", size=11)
    
    doc.add_paragraph()  # Espaço
    
    p_data_label = add_paragraph(doc, "Data:", size=11, bold=True)
    
    p_data = add_paragraph(doc, "{{ data_extenso }}", size=11)
    
    # ── ASSINATURA ──
    doc.add_paragraph()
    
    p_assinatura = add_paragraph(doc, "_" * 40, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_assinatura_nome = add_paragraph(doc, "A ARAUJO SERVIÇOS LTDA ME", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    p_assinatura_cnpj = add_paragraph(doc, "CNPJ: {{ empresa_cnpj }}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Salvar template
    template_path = "template/recibo_template.docx"
    doc.save(template_path)
    print(f"Template oficial Araujo Prev criado com sucesso em: {template_path}")

if __name__ == "__main__":
    create_araujo_receipt_template()
